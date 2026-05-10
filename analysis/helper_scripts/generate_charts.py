import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io , sys




if len(sys.argv) != 4:
    print("Usage: python script.py <count> <filter_ltp_range> <fo>")
    sys.exit(1)

n = int(sys.argv[1])
ltp_list= sys.argv[2].split('-')
start_ltp = float(ltp_list[0])
end_ltp = float(ltp_list[1])
fo = True if sys.argv[3] == 'f' else False

value_col = 'Value (₹ Lakhs) - Futures' if  fo else 'Value (₹ Lakhs) - Options (Premium)'


files = [f"data/wl{i}.csv" for i in range(1, n + 1)]
days = [f'D{i}' for i in range(1,n+1)]



# Step 2: Read and combine data from all CSV files
print("Reading CSV files...")
all_data = []
i=0
for filename in files:
    df = pd.read_csv(filename)
    # Extract only the required columns
    df_subset = df[['Symbol', value_col]].copy()
    df_subset['Day'] = days[i]
    all_data.append(df_subset)
    print(f"  {filename}: {len(df_subset)} symbols loaded")
    i+=1

# Step 3: Combine all data into a single dataframe
combined_df = pd.concat(all_data, ignore_index=True)
print(f"\nTotal records: {len(combined_df)}")
print(f"Unique symbols: {combined_df['Symbol'].nunique()}")

# Step 4: Pivot data for easy plotting (symbols as columns, days as rows)
pivot_df = combined_df.pivot(index='Day', columns='Symbol', values=value_col)
pivot_df = pivot_df.reindex(days)

print(f"\nData prepared. Shape: {pivot_df.shape}")

# Step 5: Generate plots for each symbol and store them in memory
print("\nGenerating plots...")
plot_images = []

for idx, symbol in enumerate(pivot_df.columns):
    # Create figure
    plt.figure(figsize=(10, 6))
    plt.plot(pivot_df.index, pivot_df[symbol], marker='o', linewidth=2, 
             markersize=8, color='steelblue')
    plt.title(f'{symbol}', fontsize=14, fontweight='bold')
    plt.xlabel('Day', fontsize=12)
    plt.ylabel('Vol', fontsize=12)
    plt.grid(True, alpha=0.3)
    plt.tight_layout()

    # Save plot to bytes buffer (in-memory)
    img_bytes = io.BytesIO()
    plt.savefig(img_bytes, format='png', dpi=150, bbox_inches='tight')
    img_bytes.seek(0)
    plt.close()

    # Store the symbol name and image
    plot_images.append((symbol, img_bytes))

    if (idx + 1) % 50 == 0:
        print(f"  Generated {idx + 1} plots...")

print(f"\nTotal plots generated: {len(plot_images)}")

# Step 6: Create Word document and add all plots
print("\nCreating Word document...")
doc = Document()

# Add title
title = doc.add_heading('Futures Value Analysis - All Symbols', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add introduction
intro = doc.add_paragraph()
intro.add_run('Futures Values Over 5 Days\n').bold = True
doc.add_paragraph(f'Total Symbols: {len(plot_images)}')
doc.add_paragraph('Each chart shows the trend of futures values across all 5 days')

# Add page break
doc.add_page_break()

# Add each plot to the document
for idx, (symbol, img_bytes) in enumerate(plot_images):
    # Add symbol heading
    doc.add_heading(f'{symbol}', level=2)

    # Add the image
    try:
        doc.add_picture(img_bytes, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        doc.add_paragraph(f'Error adding image for {symbol}: {str(e)}')

    # Add spacing
    doc.add_paragraph()

    # Add page break after every 5 charts (except the last)
    if (idx + 1) % 5 == 0 and idx < len(plot_images) - 1:
        doc.add_page_break()

    if (idx + 1) % 50 == 0:
        print(f"  Added {idx + 1} charts to document...")

# Save the document
output_filename = 'Futures_Value_Charts_All_Symbols.docx'
doc.save(output_filename)

print(f"\n{'='*60}")
print(f"SUCCESS!")
print(f"{'='*60}")
print(f"Document created: {output_filename}")
print(f"Total plots added: {len(plot_images)}")
print(f"{'='*60}")
